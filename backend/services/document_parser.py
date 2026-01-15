"""
Document Parser Service

Extracts SOW scope tasks from Word and PDF documents,
and parses LOE data from Excel files.
"""

import re
from pathlib import Path
from typing import List, Optional, Tuple, Dict, Any

from docx import Document
from docx.table import Table
import pdfplumber
from openpyxl import load_workbook

from models.schemas import SOWTask, LOEEntry, ColumnMappingRequest, ExcelColumn


class DocumentParser:
    """Parses SOW (Word/PDF) and LOE (Excel) documents."""
    
    # Keywords to identify scope tables
    SCOPE_TABLE_KEYWORDS = [
        "phase", "task", "deliverable", "activity", "description",
        "owner", "responsibility", "milestone", "scope"
    ]
    
    # Keywords to identify different phases
    PHASE_KEYWORDS = [
        "phase", "stage", "step", "milestone"
    ]
    
    def parse_sow_document(self, file_path: Path) -> List[SOWTask]:
        """
        Parse SOW document and extract scope tasks.
        
        Args:
            file_path: Path to the SOW document (Word or PDF)
            
        Returns:
            List of extracted SOWTask objects
        """
        suffix = file_path.suffix.lower()
        
        if suffix == ".docx":
            return self._parse_word_document(file_path)
        elif suffix == ".pdf":
            return self._parse_pdf_document(file_path)
        else:
            raise ValueError(f"Unsupported document format: {suffix}")
    
    def _parse_word_document(self, file_path: Path) -> List[SOWTask]:
        """Extract scope tasks from a Word document."""
        doc = Document(file_path)
        tasks = []
        
        # Look for tables that contain scope information
        for table in doc.tables:
            table_tasks = self._extract_tasks_from_table(table)
            if table_tasks:
                tasks.extend(table_tasks)
        
        # If no tables found, try to extract from paragraphs
        if not tasks:
            tasks = self._extract_tasks_from_paragraphs(doc)
        
        return tasks
    
    def _extract_tasks_from_table(self, table: Table) -> List[SOWTask]:
        """Extract tasks from a Word table."""
        tasks = []
        
        # Get all rows
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        
        if len(rows) < 2:
            return tasks
        
        # Find header row
        header_row = rows[0]
        header_lower = [h.lower() for h in header_row]
        
        # Check if this looks like a scope table
        if not any(kw in " ".join(header_lower) for kw in self.SCOPE_TABLE_KEYWORDS):
            return tasks
        
        # Map columns
        col_mapping = self._map_sow_columns(header_lower)
        
        if "task" not in col_mapping:
            return tasks
        
        # Extract tasks from data rows
        current_phase = "General"
        
        for row in rows[1:]:
            if len(row) <= max(col_mapping.values()):
                continue
            
            # Get phase (or use current)
            if "phase" in col_mapping:
                phase_val = row[col_mapping["phase"]].strip()
                if phase_val:
                    current_phase = phase_val
            
            # Get task
            task_val = row[col_mapping["task"]].strip()
            if not task_val:
                continue
            
            # Get description
            description = ""
            if "description" in col_mapping:
                description = row[col_mapping["description"]].strip()
            
            # Get owner
            owner = "TeraSky"
            if "owner" in col_mapping:
                owner_val = row[col_mapping["owner"]].strip()
                if owner_val:
                    owner = owner_val
            
            tasks.append(SOWTask(
                phase=current_phase,
                task=task_val,
                description=description or task_val,
                owner=owner,
            ))
        
        return tasks
    
    def _map_sow_columns(self, headers: List[str]) -> Dict[str, int]:
        """Map header names to column indices."""
        mapping = {}
        
        for idx, header in enumerate(headers):
            header_clean = header.lower().strip()
            
            if any(kw in header_clean for kw in ["phase", "stage"]):
                mapping["phase"] = idx
            elif any(kw in header_clean for kw in ["task", "activity", "deliverable", "scope"]):
                if "task" not in mapping:  # Prefer first match
                    mapping["task"] = idx
            elif any(kw in header_clean for kw in ["description", "detail", "note"]):
                mapping["description"] = idx
            elif any(kw in header_clean for kw in ["owner", "responsible", "assigned"]):
                mapping["owner"] = idx
        
        return mapping
    
    def _extract_tasks_from_paragraphs(self, doc: Document) -> List[SOWTask]:
        """Extract tasks from document paragraphs (fallback method)."""
        tasks = []
        current_phase = "General"
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Check if this is a heading (potential phase)
            if para.style and "Heading" in para.style.name:
                if any(kw in text.lower() for kw in ["phase", "stage", "section"]):
                    current_phase = text
                continue
            
            # Check for numbered/bulleted items that might be tasks
            if re.match(r'^[\d.]+\s+', text) or text.startswith(('•', '-', '*')):
                task_text = re.sub(r'^[\d.•\-*]+\s*', '', text)
                if len(task_text) > 10:  # Reasonable task length
                    tasks.append(SOWTask(
                        phase=current_phase,
                        task=task_text[:100],  # Truncate long titles
                        description=task_text,
                        owner="TeraSky",
                    ))
        
        return tasks
    
    def _parse_pdf_document(self, file_path: Path) -> List[SOWTask]:
        """Extract scope tasks from a PDF document."""
        tasks = []
        
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # Try to extract tables
                tables = page.extract_tables()
                
                for table in tables:
                    if not table or len(table) < 2:
                        continue
                    
                    table_tasks = self._extract_tasks_from_pdf_table(table)
                    if table_tasks:
                        tasks.extend(table_tasks)
                
                # If no tables, try text extraction
                if not tasks:
                    text = page.extract_text()
                    if text:
                        text_tasks = self._extract_tasks_from_text(text)
                        tasks.extend(text_tasks)
        
        return tasks
    
    def _extract_tasks_from_pdf_table(self, table: List[List[str]]) -> List[SOWTask]:
        """Extract tasks from a PDF table."""
        tasks = []
        
        # Clean table data
        table = [[str(cell).strip() if cell else "" for cell in row] for row in table]
        
        if len(table) < 2:
            return tasks
        
        # Find header row
        header_row = table[0]
        header_lower = [h.lower() for h in header_row]
        
        # Check if this looks like a scope table
        if not any(kw in " ".join(header_lower) for kw in self.SCOPE_TABLE_KEYWORDS):
            return tasks
        
        # Map columns
        col_mapping = self._map_sow_columns(header_lower)
        
        if "task" not in col_mapping:
            return tasks
        
        # Extract tasks
        current_phase = "General"
        
        for row in table[1:]:
            if len(row) <= max(col_mapping.values()):
                continue
            
            if "phase" in col_mapping:
                phase_val = row[col_mapping["phase"]].strip()
                if phase_val:
                    current_phase = phase_val
            
            task_val = row[col_mapping["task"]].strip()
            if not task_val:
                continue
            
            description = ""
            if "description" in col_mapping:
                description = row[col_mapping["description"]].strip()
            
            owner = "TeraSky"
            if "owner" in col_mapping:
                owner_val = row[col_mapping["owner"]].strip()
                if owner_val:
                    owner = owner_val
            
            tasks.append(SOWTask(
                phase=current_phase,
                task=task_val,
                description=description or task_val,
                owner=owner,
            ))
        
        return tasks
    
    def _extract_tasks_from_text(self, text: str) -> List[SOWTask]:
        """Extract tasks from plain text (fallback)."""
        tasks = []
        lines = text.split('\n')
        current_phase = "General"
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for phase headers
            if re.match(r'^(phase|stage|section)\s*[\d:]+', line.lower()):
                current_phase = line
                continue
            
            # Check for numbered items
            if re.match(r'^[\d.]+\s+', line):
                task_text = re.sub(r'^[\d.]+\s*', '', line)
                if len(task_text) > 10:
                    tasks.append(SOWTask(
                        phase=current_phase,
                        task=task_text[:100],
                        description=task_text,
                        owner="TeraSky",
                    ))
        
        return tasks
    
    def parse_loe_excel(
        self,
        file_path: Path,
        column_mapping: ColumnMappingRequest,
        sheet_name: Optional[str] = None,
    ) -> List[LOEEntry]:
        """
        Parse LOE entries from an Excel file.
        
        Args:
            file_path: Path to the Excel file
            column_mapping: Configuration for which columns to read
            sheet_name: Name of the sheet to read (defaults to first sheet)
            
        Returns:
            List of LOEEntry objects
        """
        workbook = load_workbook(filename=file_path, read_only=True, data_only=True)
        
        # Get the worksheet
        if sheet_name:
            if sheet_name not in workbook.sheetnames:
                raise ValueError(f"Sheet '{sheet_name}' not found. Available: {workbook.sheetnames}")
            worksheet = workbook[sheet_name]
        else:
            worksheet = workbook.active
        
        entries = self._parse_loe_worksheet(worksheet, column_mapping)
        workbook.close()
        
        return entries
    
    def _parse_loe_worksheet(
        self,
        worksheet,
        mapping: ColumnMappingRequest,
    ) -> List[LOEEntry]:
        """Parse entries from a worksheet using the column mapping."""
        # Check if we're using index-based column references (e.g., "Column 4")
        is_index_based = bool(re.match(r'^Column\s+\d+$', mapping.task_column.strip(), re.IGNORECASE))
        
        header_row = None
        col_indices = {}
        
        if is_index_based:
            # For index-based columns, find the first non-empty row as header
            for row_idx in range(1, 11):
                row_values = [
                    str(cell.value).strip() if cell.value else ""
                    for cell in worksheet[row_idx]
                ]
                if any(v for v in row_values):  # Non-empty row
                    header_row = row_idx
                    col_indices = self._map_loe_columns(row_values, mapping)
                    break
        else:
            # Standard name-based column search
            for row_idx in range(1, 11):
                row_values = [
                    str(cell.value).strip() if cell.value else ""
                    for cell in worksheet[row_idx]
                ]
                
                # Check if this row contains the required columns
                if self._find_column(row_values, mapping.task_column) is not None:
                    header_row = row_idx
                    col_indices = self._map_loe_columns(row_values, mapping)
                    break
        
        if header_row is None:
            raise ValueError(f"Could not find header row with column '{mapping.task_column}'")
        
        # Validate required columns were found
        if "task" not in col_indices:
            raise ValueError(f"Required column '{mapping.task_column}' not found")
        if "days" not in col_indices:
            raise ValueError(f"Required column '{mapping.days_column}' not found")
        
        # Parse data rows
        entries = []
        for row_idx in range(header_row + 1, worksheet.max_row + 1):
            row_data = list(worksheet[row_idx])
            
            # Get task value
            task_value = self._get_cell_value(row_data, col_indices.get("task"))
            if not task_value:
                continue
            
            # Get days value
            days_value = self._get_numeric_value(row_data, col_indices.get("days"))
            if days_value is None:
                continue
            
            # Get optional values
            phase_value = self._get_cell_value(row_data, col_indices.get("phase"))
            risk_value = self._get_numeric_value(row_data, col_indices.get("risk"))
            total_value = self._get_numeric_value(row_data, col_indices.get("total"))
            
            entry = LOEEntry(
                task=task_value,
                phase=phase_value,
                days=days_value,
                risk_buffer=risk_value,
                total_days=total_value,
            )
            entries.append(entry)
        
        return entries
    
    def _find_column(self, row_values: List[str], column_name: str) -> Optional[int]:
        """
        Find column index by name (case-insensitive).
        
        Also handles placeholder column names like "Column 4" which refer
        to columns by their index (1-based) when headers are empty.
        """
        column_name_stripped = column_name.strip()
        
        # Check if this is an index-based column reference (e.g., "Column 4")
        index_match = re.match(r'^Column\s+(\d+)$', column_name_stripped, re.IGNORECASE)
        if index_match:
            col_index = int(index_match.group(1)) - 1  # Convert to 0-based
            if 0 <= col_index < len(row_values):
                return col_index
            return None
        
        # Standard name-based lookup
        column_name_lower = column_name_stripped.lower()
        for idx, value in enumerate(row_values):
            if value.lower().strip() == column_name_lower:
                return idx
        return None
    
    def _map_loe_columns(
        self,
        row_values: List[str],
        mapping: ColumnMappingRequest,
    ) -> Dict[str, int]:
        """Map column names to indices."""
        result = {}
        
        task_idx = self._find_column(row_values, mapping.task_column)
        if task_idx is not None:
            result["task"] = task_idx
        
        days_idx = self._find_column(row_values, mapping.days_column)
        if days_idx is not None:
            result["days"] = days_idx
        
        if mapping.phase_column:
            phase_idx = self._find_column(row_values, mapping.phase_column)
            if phase_idx is not None:
                result["phase"] = phase_idx
        
        if mapping.risk_column:
            risk_idx = self._find_column(row_values, mapping.risk_column)
            if risk_idx is not None:
                result["risk"] = risk_idx
        
        if mapping.total_column:
            total_idx = self._find_column(row_values, mapping.total_column)
            if total_idx is not None:
                result["total"] = total_idx
        
        return result
    
    def _get_cell_value(self, row_data: list, col_idx: Optional[int]) -> Optional[str]:
        """Get string value from cell."""
        if col_idx is None or col_idx >= len(row_data):
            return None
        value = row_data[col_idx].value
        if value is None:
            return None
        return str(value).strip()
    
    def _get_numeric_value(self, row_data: list, col_idx: Optional[int]) -> Optional[float]:
        """Get numeric value from cell."""
        if col_idx is None or col_idx >= len(row_data):
            return None
        value = row_data[col_idx].value
        if value is None:
            return None
        try:
            return float(value)
        except (ValueError, TypeError):
            return None
    
    def preview_excel(self, file_path: Path) -> Tuple[List[str], List[ExcelColumn], int]:
        """
        Preview Excel file structure for column mapping UI.
        
        Returns:
            Tuple of (sheet names, columns with samples, row count)
        """
        workbook = load_workbook(filename=file_path, read_only=True, data_only=True)
        sheets = workbook.sheetnames
        
        worksheet = workbook.active
        
        # Find header row
        header_row = None
        for row_idx in range(1, 11):
            row_values = [
                str(cell.value).strip() if cell.value else ""
                for cell in worksheet[row_idx]
            ]
            if any(v for v in row_values):  # Non-empty row
                header_row = row_idx
                break
        
        if header_row is None:
            workbook.close()
            return sheets, [], 0
        
        # Get headers
        headers = [
            str(cell.value).strip() if cell.value else f"Column {i+1}"
            for i, cell in enumerate(worksheet[header_row])
        ]
        
        # Get sample values (next 3 rows)
        columns = []
        for col_idx, header in enumerate(headers):
            samples = []
            for row_idx in range(header_row + 1, min(header_row + 4, worksheet.max_row + 1)):
                row = list(worksheet[row_idx])
                if col_idx < len(row) and row[col_idx].value is not None:
                    samples.append(str(row[col_idx].value)[:50])  # Truncate long values
            
            columns.append(ExcelColumn(name=header, sample_values=samples))
        
        row_count = worksheet.max_row - header_row
        workbook.close()
        
        return sheets, columns, row_count
