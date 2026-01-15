"""
Validation Router

Handles file uploads and SOW vs LOE validation.
"""

import uuid
import shutil
from datetime import datetime
from pathlib import Path
from typing import Optional

from fastapi import APIRouter, File, UploadFile, HTTPException, Form
from fastapi.responses import FileResponse

from models.schemas import (
    ColumnMappingRequest,
    ValidationRequest,
    ValidationResponse,
    UploadResponse,
    ExcelPreview,
    SOWTask,
)
from services.document_parser import DocumentParser
from services.validator_service import ValidatorService


router = APIRouter()

# Storage directories
UPLOAD_DIR = Path(__file__).parent.parent / "uploads"
REPORTS_DIR = Path(__file__).parent.parent / "reports"
UPLOAD_DIR.mkdir(exist_ok=True)
REPORTS_DIR.mkdir(exist_ok=True)

# In-memory file registry (in production, use a database)
_file_registry: dict = {}


@router.post("/upload", response_model=UploadResponse)
async def upload_file(file: UploadFile = File(...)):
    """
    Upload a SOW (docx/pdf) or LOE (xlsx) file.
    
    Returns a file ID that can be used for validation.
    """
    # Validate file type
    allowed_extensions = {".docx", ".pdf", ".xlsx", ".xls"}
    suffix = Path(file.filename).suffix.lower()
    
    if suffix not in allowed_extensions:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {suffix}. Allowed: {', '.join(allowed_extensions)}"
        )
    
    # Generate unique file ID
    file_id = str(uuid.uuid4())
    
    # Determine file type
    if suffix in {".docx", ".pdf"}:
        file_type = "sow"
    else:
        file_type = "loe"
    
    # Save file
    file_path = UPLOAD_DIR / f"{file_id}{suffix}"
    
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to save file: {str(e)}"
        )
    
    # Get file size
    file_size = file_path.stat().st_size
    
    # Register file
    _file_registry[file_id] = {
        "filename": file.filename,
        "path": str(file_path),
        "type": file_type,
        "uploaded_at": datetime.now().isoformat(),
    }
    
    return UploadResponse(
        file_id=file_id,
        filename=file.filename,
        file_type=file_type,
        size_bytes=file_size,
    )


@router.get("/preview-excel/{file_id}", response_model=ExcelPreview)
async def preview_excel(file_id: str):
    """
    Preview Excel file structure for column mapping configuration.
    
    Returns sheet names, column headers, and sample values.
    """
    if file_id not in _file_registry:
        raise HTTPException(status_code=404, detail="File not found")
    
    file_info = _file_registry[file_id]
    file_path = Path(file_info["path"])
    
    if not file_path.suffix.lower() in {".xlsx", ".xls"}:
        raise HTTPException(status_code=400, detail="Not an Excel file")
    
    parser = DocumentParser()
    
    try:
        sheets, columns, row_count = parser.preview_excel(file_path)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to preview Excel file: {str(e)}"
        )
    
    return ExcelPreview(
        file_id=file_id,
        sheets=sheets,
        columns=columns,
        row_count=row_count,
    )


@router.post("/validate", response_model=ValidationResponse)
async def validate_sow_loe(request: ValidationRequest):
    """
    Validate SOW document against LOE Excel file.
    
    Performs task matching, complexity analysis, and duration validation.
    """
    # Get file info
    if request.sow_file_id not in _file_registry:
        raise HTTPException(status_code=404, detail="SOW file not found")
    if request.loe_file_id not in _file_registry:
        raise HTTPException(status_code=404, detail="LOE file not found")
    
    sow_info = _file_registry[request.sow_file_id]
    loe_info = _file_registry[request.loe_file_id]
    
    sow_path = Path(sow_info["path"])
    loe_path = Path(loe_info["path"])
    
    parser = DocumentParser()
    validator = ValidatorService()
    
    # Parse SOW document
    try:
        sow_tasks = parser.parse_sow_document(sow_path)
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to parse SOW document: {str(e)}"
        )
    
    if not sow_tasks:
        raise HTTPException(
            status_code=400,
            detail="No tasks found in SOW document. Ensure the document contains a scope table."
        )
    
    # Parse LOE Excel
    try:
        loe_entries = parser.parse_loe_excel(
            loe_path,
            request.column_mapping,
            request.sheet_name,
        )
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to parse LOE Excel: {str(e)}"
        )
    
    if not loe_entries:
        raise HTTPException(
            status_code=400,
            detail="No entries found in LOE Excel. Check column mapping configuration."
        )
    
    # Perform validation
    result = validator.validate(
        sow_tasks=sow_tasks,
        loe_entries=loe_entries,
        customer_name=request.customer_name,
        project_name=request.project_name,
    )
    
    return result


@router.post("/validate-json", response_model=ValidationResponse)
async def validate_sow_json(
    sow_tasks: list[dict],
    loe_file_id: str,
    column_mapping: ColumnMappingRequest,
    sheet_name: Optional[str] = None,
    customer_name: str = "Customer",
    project_name: str = "Project",
):
    """
    Validate SOW tasks (as JSON) against LOE Excel file.
    
    Use this if SOW tasks are already extracted or manually entered.
    """
    if loe_file_id not in _file_registry:
        raise HTTPException(status_code=404, detail="LOE file not found")
    
    loe_info = _file_registry[loe_file_id]
    loe_path = Path(loe_info["path"])
    
    parser = DocumentParser()
    validator = ValidatorService()
    
    # Convert JSON to SOWTask objects
    try:
        tasks = [SOWTask(**task) for task in sow_tasks]
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid SOW task format: {str(e)}"
        )
    
    # Parse LOE Excel
    try:
        loe_entries = parser.parse_loe_excel(loe_path, column_mapping, sheet_name)
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to parse LOE Excel: {str(e)}"
        )
    
    # Perform validation
    result = validator.validate(
        sow_tasks=tasks,
        loe_entries=loe_entries,
        customer_name=customer_name,
        project_name=project_name,
    )
    
    return result


@router.post("/generate-report/{validation_id}")
async def generate_report(
    validation_id: str,
    validation_result: dict,
    format: str = "docx",
):
    """
    Generate a downloadable report from validation results.
    
    Supported formats: docx, pdf (future)
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Create Word document
    doc = Document()
    
    # Title
    title = doc.add_heading("SOW vs LOE Validation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Customer/Project
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        f"{validation_result.get('customer_name', 'Customer')} - "
        f"{validation_result.get('project_name', 'Project')}"
    )
    run.bold = True
    run.font.size = Pt(14)
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(datetime.now().strftime("%B %d, %Y")).italic = True
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading("Executive Summary", 1)
    status = validation_result.get('status', 'UNKNOWN')
    summary_para = doc.add_paragraph()
    summary_para.add_run("Validation Status: ").bold = True
    status_run = summary_para.add_run(status)
    status_run.bold = True
    
    # Summary Statistics
    doc.add_heading("Summary", 2)
    stats_table = doc.add_table(rows=5, cols=2)
    stats_table.style = "Table Grid"
    
    stats = [
        ("SOW Tasks", str(validation_result.get('total_sow_tasks', 0))),
        ("LOE Entries", str(validation_result.get('total_loe_entries', 0))),
        ("Matched Tasks", str(validation_result.get('matched_tasks', 0))),
        ("Total LOE Days", f"{validation_result.get('total_loe_days', 0):.1f}"),
        ("Variance", f"{validation_result.get('total_variance_percent', 0):.1f}%"),
    ]
    
    for i, (label, value) in enumerate(stats):
        stats_table.rows[i].cells[0].text = label
        stats_table.rows[i].cells[1].text = value
    
    # Issues
    issues = validation_result.get('critical_issues', [])
    if issues:
        doc.add_heading("Critical Issues", 2)
        for issue in issues:
            doc.add_paragraph(issue, style="List Bullet")
    
    # Warnings
    warnings = validation_result.get('warnings', [])
    if warnings:
        doc.add_heading("Warnings", 2)
        for warning in warnings:
            doc.add_paragraph(warning, style="List Bullet")
    
    # Recommendations
    recommendations = validation_result.get('recommendations', [])
    if recommendations:
        doc.add_heading("Recommendations", 1)
        for rec in recommendations:
            doc.add_paragraph(rec, style="List Bullet")
    
    # Save document
    filename = f"Validation_Report_{validation_id}.docx"
    output_path = REPORTS_DIR / filename
    doc.save(output_path)
    
    return {
        "status": "success",
        "filename": filename,
        "download_url": f"/reports/{filename}",
    }


@router.get("/download-report/{filename}")
async def download_report(filename: str):
    """Download a generated report."""
    file_path = REPORTS_DIR / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Report not found")
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@router.get("/files")
async def list_files():
    """List all uploaded files."""
    return {
        "files": [
            {
                "file_id": fid,
                "filename": info["filename"],
                "type": info["type"],
                "uploaded_at": info["uploaded_at"],
            }
            for fid, info in _file_registry.items()
        ]
    }


@router.delete("/files/{file_id}")
async def delete_file(file_id: str):
    """Delete an uploaded file."""
    if file_id not in _file_registry:
        raise HTTPException(status_code=404, detail="File not found")
    
    file_info = _file_registry[file_id]
    file_path = Path(file_info["path"])
    
    if file_path.exists():
        file_path.unlink()
    
    del _file_registry[file_id]
    
    return {"status": "deleted", "file_id": file_id}
